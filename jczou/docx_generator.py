import json
import sys
import os
import base64
import tempfile
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_run_with_font(paragraph, text, size_pt=10, bold=False, color_rgb=None, font_name='Microsoft YaHei'):
    run = paragraph.add_run(text)
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color_rgb:
        run.font.color.rgb = color_rgb
    
    # Set font XML descriptors for compatibility with both English and Chinese
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.append(rFonts)
    return run

def add_section_title(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = False
    p.paragraph_format.keep_together = False
    
    # Bottom border for section title
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12') # 1.5 pt
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '000000') # Black
    pBdr.append(bottom)
    pPr.append(pBdr)
    
    add_run_with_font(p, text, size_pt=12, bold=True, color_rgb=RGBColor(0, 0, 0))
    return p

def main():
    if len(sys.argv) < 3:
        print("Usage: docx_generator.py <input_json_path> <output_docx_path>")
        sys.exit(1)
        
    input_json = sys.argv[1]
    output_docx = sys.argv[2]
    
    if not os.path.exists(input_json):
        print(f"Error: Input JSON file not found: {input_json}")
        sys.exit(1)
        
    with open(input_json, "r", encoding="utf-8") as f:
        pkg = json.load(f)
        
    doc = Document()
    
    # Set page margins (A4: 210 x 297 mm, 1 inch = 72 pt)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.page_width = Inches(8.27)  # A4
        section.page_height = Inches(11.69)
        
    # Set default style font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(10)
    
    # Check if photo base64 is present
    photo_base64 = pkg.get("photoBase64", "")
    temp_img_path = None
    
    if photo_base64:
        try:
            img_data = base64.b64decode(photo_base64)
            # Create a temporary image file
            fd, temp_img_path = tempfile.mkstemp(suffix=".png")
            os.close(fd)
            with open(temp_img_path, "wb") as f_img:
                f_img.write(img_data)
        except Exception as e:
            print(f"Warning: Failed to decode photoBase64: {e}")
            temp_img_path = None

    if temp_img_path:
        # Header layout with 2-column table: Left for text, Right for photo
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.allow_autofit = False
        
        # Set cell widths
        # Page body width is 7.27 inches. Left: 5.77, Right: 1.50
        table.columns[0].width = Inches(5.77)
        table.columns[1].width = Inches(1.50)
        
        cell_text = table.cell(0, 0)
        cell_photo = table.cell(0, 1)
        
        # Remove cell padding/margins if necessary, but default table style is fine
        p_name = cell_text.paragraphs[0]
        p_name.paragraph_format.space_after = Pt(2)
        add_run_with_font(p_name, pkg.get("fullName", "未填写"), size_pt=18, bold=True)
        
        p_contact = cell_text.add_paragraph()
        p_contact.paragraph_format.space_after = Pt(2)
        add_run_with_font(p_contact, pkg.get("contact", ""), size_pt=9, color_rgb=RGBColor(0, 0, 0))
        
        target_job = pkg.get("targetJob", "")
        if target_job:
            p_target = cell_text.add_paragraph()
            p_target.paragraph_format.space_after = Pt(8)
            add_run_with_font(p_target, f"求职意向：{target_job}", size_pt=9.5, bold=True, color_rgb=RGBColor(0, 0, 0))
            
        # Add photo to right cell, right aligned
        p_photo = cell_photo.paragraphs[0]
        p_photo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_photo.paragraph_format.space_after = Pt(0)
        try:
            # Set photo size to 1.0 inch width, height adjusts proportionally (usually 1.3 inches)
            p_photo.add_run().add_picture(temp_img_path, width=Inches(1.0))
        except Exception as e:
            print(f"Warning: Failed to add photo to document: {e}")
    else:
        # Standard centered header without photo
        p_name = doc.add_paragraph()
        p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_name.paragraph_format.space_after = Pt(2)
        add_run_with_font(p_name, pkg.get("fullName", "未填写"), size_pt=18, bold=True)
        
        p_contact = doc.add_paragraph()
        p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_contact.paragraph_format.space_after = Pt(2)
        add_run_with_font(p_contact, pkg.get("contact", ""), size_pt=9, color_rgb=RGBColor(0, 0, 0))
        
        target_job = pkg.get("targetJob", "")
        if target_job:
            p_target = doc.add_paragraph()
            p_target.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_target.paragraph_format.space_after = Pt(8)
            add_run_with_font(p_target, f"求职意向：{target_job}", size_pt=9.5, bold=True, color_rgb=RGBColor(0, 0, 0))
            
    # 2. Education
    add_section_title(doc, "教育背景")
    
    # GPA
    p_gpa = doc.add_paragraph()
    p_gpa.paragraph_format.space_after = Pt(4)
    run_gpa_lbl = add_run_with_font(p_gpa, "综合 GPA：", bold=True)
    gpa_val = pkg.get("finalGpa", 0.0)
    gpa4 = gpa_val / 100.0 * 4.0
    add_run_with_font(p_gpa, f"{gpa_val:.1f} / 100（{gpa4:.2f} / 4.0）", bold=True, color_rgb=RGBColor(0, 0, 0))
    
    # Courses (Inline list)
    courses = pkg.get("courses", [])
    if courses:
        p_courses = doc.add_paragraph()
        p_courses.paragraph_format.space_after = Pt(6)
        p_courses.paragraph_format.line_spacing = 1.15
        add_run_with_font(p_courses, "核心课程：", bold=True)
        course_str_list = [f"{c['name']} {c['grade']:.1f}" for c in courses]
        add_run_with_font(p_courses, ", ".join(course_str_list))
        
    # 3. Experiences
    exp_groups = {
        "internship": "实习经历",
        "project": "项目经历",
        "activity": "校园经历",
        "other": "其他经历"
    }
    
    # Sort experiences by type order
    exps = pkg.get("experiences", [])
    for etype, title_cn in exp_groups.items():
        group = [e for e in exps if e.get("type") == etype]
        if not group:
            continue
            
        add_section_title(doc, title_cn)
        
        for idx, e in enumerate(group):
            p_header = doc.add_paragraph()
            p_header.paragraph_format.space_after = Pt(2)
            p_header.paragraph_format.space_before = Pt(4) if idx > 0 else Pt(0)
            
            # Align dates to the right using tab stop (A4 content width is about 7.27 inches)
            p_header.paragraph_format.tab_stops.add_tab_stop(Inches(7.27), WD_TAB_ALIGNMENT.RIGHT)
            
            add_run_with_font(p_header, e.get("title", ""), bold=True)
            add_run_with_font(p_header, f"     {e.get('organization', '')}", color_rgb=RGBColor(0, 0, 0))
            
            # Use computed date range from C++ if passed, or compute here
            date_range = e.get("dateRange", "")
            if not date_range:
                # Formatting QDate ISO format to YYYY.MM
                def fmt_date(d_str):
                    if not d_str:
                        return ""
                    parts = d_str.split('-')
                    if len(parts) >= 2:
                        return f"{parts[0]}.{parts[1]}"
                    return d_str
                    
                s_date = fmt_date(e.get("startDate", ""))
                e_date = fmt_date(e.get("endDate", ""))
                if s_date and e_date:
                    date_range = f"{s_date} - {e_date}"
                elif s_date:
                    date_range = f"{s_date} - 至今"
                elif e_date:
                    date_range = e_date
                    
            p_header.add_run(f"\t{date_range}")
            
            desc = e.get("description", "")
            if desc:
                p_desc = doc.add_paragraph()
                p_desc.paragraph_format.space_after = Pt(4)
                p_desc.paragraph_format.line_spacing = 1.15
                run_desc = add_run_with_font(p_desc, desc, size_pt=9.5, color_rgb=RGBColor(0, 0, 0))
                
    # 4. Skills / Awards
    skills = pkg.get("skills", "")
    if skills:
        add_section_title(doc, "技能 / 证书")
        p_skills = doc.add_paragraph()
        p_skills.paragraph_format.space_after = Pt(4)
        p_skills.paragraph_format.line_spacing = 1.2
        add_run_with_font(p_skills, skills, size_pt=9.5)
        
    doc.save(output_docx)
    print("Docx generated successfully")
    
    # Cleanup temp image file
    if temp_img_path and os.path.exists(temp_img_path):
        try:
            os.remove(temp_img_path)
        except Exception as e:
            print(f"Warning: Failed to remove temporary photo file: {e}")

if __name__ == "__main__":
    main()
